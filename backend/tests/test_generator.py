from io import BytesIO

from docx import Document
from fastapi.testclient import TestClient

from app.generator import extract_document_text
from app.main import app


def test_extracts_text_and_word_documents():
    assert extract_document_text("requirements.md", b"# Login\nUsers can sign in.") == "# Login\nUsers can sign in."

    stream = BytesIO()
    document = Document()
    document.add_heading("Checkout", level=1)
    document.add_paragraph("A customer can pay by card.")
    document.save(stream)
    assert extract_document_text("requirements.docx", stream.getvalue()) == "Checkout\nA customer can pay by card."


def test_generation_rejects_unsupported_documents_and_requires_api_key(monkeypatch):
    client = TestClient(app)
    unsupported = client.post("/api/generation/cases", files={"file": ("requirements.csv", b"a,b", "text/csv")})
    assert unsupported.status_code == 400

    monkeypatch.delenv("DEEPSEEK_API_KEY", raising=False)
    missing_key = client.post("/api/generation/cases", files={"file": ("requirements.txt", b"Users can sign in with valid credentials.", "text/plain")})
    assert missing_key.status_code == 503
    assert "API key" in missing_key.json()["detail"]
