from __future__ import annotations

from io import BytesIO
import json
import os
from pathlib import Path

import httpx
from docx import Document
from pypdf import PdfReader

from .llm import DEEPSEEK_BASE_URL, DEEPSEEK_MODEL
from .models import GeneratedCase, GenerationResponse


MAX_DOCUMENT_CHARS = 45_000


def extract_document_text(filename: str, content: bytes) -> str:
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        text = "\n".join(page.extract_text() or "" for page in PdfReader(BytesIO(content)).pages)
    elif suffix == ".docx":
        document = Document(BytesIO(content))
        parts = [paragraph.text for paragraph in document.paragraphs]
        parts.extend(" | ".join(cell.text for cell in row.cells) for table in document.tables for row in table.rows)
        text = "\n".join(parts)
    elif suffix in {".txt", ".md"}:
        text = content.decode("utf-8-sig", errors="replace")
    else:
        raise ValueError("Supported document types are PDF, DOCX, Markdown and plain text.")
    text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
    if not text:
        raise ValueError("No readable text was found in the uploaded document.")
    return text[:MAX_DOCUMENT_CHARS]


async def generate_cases(filename: str, content: bytes, request_api_key: str | None = None) -> GenerationResponse:
    text = extract_document_text(filename, content)
    api_key = (request_api_key or os.getenv("DEEPSEEK_API_KEY", "")).strip()
    if not api_key:
        raise RuntimeError("Add and verify a DeepSeek API key in Project settings before generating cases.")
    payload = {
        "model": DEEPSEEK_MODEL,
        "temperature": 0.2,
        "response_format": {"type": "json_object"},
        "thinking": {"type": "disabled"},
        "messages": [
            {"role": "system", "content": "You are a senior QA analyst. Turn the supplied product document into a compact, non-duplicative test suite. Cover happy paths, validation, boundaries and failure handling that are supported by the document. Return JSON only: summary (string), cases (array of 4-20 objects). Every case must contain title, case_type (Web/API/Mobile), priority (P0/P1/P2), preconditions, numbered newline-separated test_steps, expected_result, and requirement (a short source requirement label). Do not invent product behavior or credentials."},
            {"role": "user", "content": json.dumps({"filename": filename, "document": text}, ensure_ascii=False)},
        ],
    }
    async with httpx.AsyncClient(timeout=60) as client:
        response = await client.post(f"{DEEPSEEK_BASE_URL}/chat/completions", headers={"Authorization": f"Bearer {api_key}"}, json=payload)
    if response.status_code == 401:
        raise RuntimeError("DeepSeek rejected the configured API key.")
    if not response.is_success:
        raise RuntimeError("DeepSeek could not generate test cases right now.")
    try:
        result = json.loads(response.json().get("choices", [{}])[0].get("message", {}).get("content", "{}"))
        cases = [GeneratedCase.model_validate(item) for item in result.get("cases", [])]
    except (ValueError, TypeError, IndexError) as error:
        raise RuntimeError("The AI service returned an invalid test suite.") from error
    if not cases:
        raise RuntimeError("No testable requirements were found in this document.")
    return GenerationResponse(filename=filename, summary=str(result.get("summary") or f"Generated {len(cases)} cases from {filename}."), cases=cases)
